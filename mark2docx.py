#!/usr/bin/env python3
"""
Markdown to Word Document Converter

This script converts markdown files to Microsoft Word documents (.docx).
It preserves formatting including headings, bold, italic, lists, links, 
code blocks, and tables.

Requirements:
    pip install python-docx markdown2 beautifulsoup4
"""

import argparse
import sys
from pathlib import Path
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
from bs4 import BeautifulSoup, NavigableString


class MarkdownToWordConverter:
    """Converts markdown content to a Word document."""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """Set up custom styles for the document."""
        # Define heading styles
        styles = self.doc.styles
        
        # Customize heading styles
        for i in range(1, 7):
            heading_style = styles[f'Heading {i}']
            heading_style.font.size = Pt(26 - (i * 2))
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            
        # Code style
        if 'Code' not in styles:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.CHARACTER)
            code_style.font.name = 'Courier New'
            code_style.font.size = Pt(10)
            code_style.font.color.rgb = RGBColor(0, 0, 139)
            
    def convert_markdown_to_word(self, markdown_content, output_path):
        """Convert markdown content to a Word document."""
        # Convert markdown to HTML
        html_content = markdown2.markdown(
            markdown_content,
            extras=['tables', 'fenced-code-blocks', 'code-friendly']
        )
        
        # Parse HTML and build Word document
        soup = BeautifulSoup(html_content, 'html.parser')
        self._process_element(soup)
        
        # Save the document
        self.doc.save(output_path)
        
    def _process_element(self, element, paragraph=None):
        """Recursively process HTML elements and convert to Word."""
        if isinstance(element, NavigableString):
            if paragraph is not None and str(element).strip():
                paragraph.add_run(str(element))
            return
            
        if element.name == 'h1':
            self.doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            self.doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            self.doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            self.doc.add_heading(element.get_text(), level=4)
        elif element.name == 'h5':
            self.doc.add_heading(element.get_text(), level=5)
        elif element.name == 'h6':
            self.doc.add_heading(element.get_text(), level=6)
        elif element.name == 'p':
            p = self.doc.add_paragraph()
            self._process_inline_elements(element, p)
        elif element.name == 'ul':
            self._process_list(element, ordered=False)
        elif element.name == 'ol':
            self._process_list(element, ordered=True)
        elif element.name == 'pre':
            # Code block
            code_text = element.get_text()
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 139)
        elif element.name == 'blockquote':
            # Blockquote
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(36)
            self._process_inline_elements(element, p)
        elif element.name == 'table':
            self._process_table(element)
        elif element.name == 'hr':
            # Horizontal rule
            p = self.doc.add_paragraph('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Process children for other elements
            for child in element.children:
                self._process_element(child, paragraph)
                
    def _process_inline_elements(self, element, paragraph):
        """Process inline elements like bold, italic, links, etc."""
        for child in element.children:
            if isinstance(child, NavigableString):
                paragraph.add_run(str(child))
            elif child.name == 'strong' or child.name == 'b':
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name == 'em' or child.name == 'i':
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == 'code':
                run = paragraph.add_run(child.get_text())
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 139)
            elif child.name == 'a':
                # Links - add the text and URL in parentheses
                link_text = child.get_text()
                href = child.get('href', '')
                run = paragraph.add_run(f"{link_text} ({href})")
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
            else:
                self._process_inline_elements(child, paragraph)
                
    def _process_list(self, list_element, ordered=False):
        """Process unordered or ordered lists."""
        items = list_element.find_all('li', recursive=False)
        for i, item in enumerate(items, 1):
            if ordered:
                prefix = f"{i}. "
            else:
                prefix = "• "
                
            p = self.doc.add_paragraph(prefix)
            self._process_inline_elements(item, p)
            
    def _process_table(self, table_element):
        """Process HTML tables."""
        rows = table_element.find_all('tr')
        if not rows:
            return
            
        # Count columns
        first_row = rows[0]
        cols = len(first_row.find_all(['th', 'td']))
        
        # Create table
        table = self.doc.add_table(rows=0, cols=cols)
        table.style = 'Table Grid'
        
        # Add rows
        for row in rows:
            cells = row.find_all(['th', 'td'])
            table_row = table.add_row()
            
            for i, cell in enumerate(cells):
                table_row.cells[i].text = cell.get_text().strip()
                # Make header cells bold
                if cell.name == 'th':
                    for paragraph in table_row.cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def main():
    """Main function to handle command line arguments."""
    parser = argparse.ArgumentParser(
        description='Convert Markdown files to Word documents'
    )
    parser.add_argument(
        'input',
        type=str,
        help='Input markdown file path'
    )
    parser.add_argument(
        '-o', '--output',
        type=str,
        help='Output Word document path (default: input_file.docx)'
    )
    
    args = parser.parse_args()
    
    # Validate input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file '{args.input}' not found.")
        sys.exit(1)
        
    if not input_path.is_file():
        print(f"Error: '{args.input}' is not a file.")
        sys.exit(1)
        
    # Determine output path
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_suffix('.docx')
        
    # Read markdown content
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            markdown_content = f.read()
    except Exception as e:
        print(f"Error reading file: {e}")
        sys.exit(1)
        
    # Convert to Word
    try:
        converter = MarkdownToWordConverter()
        converter.convert_markdown_to_word(markdown_content, str(output_path))
        print(f"Successfully converted '{args.input}' to '{output_path}'")
    except Exception as e:
        print(f"Error during conversion: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()